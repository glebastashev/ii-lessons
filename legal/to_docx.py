# -*- coding: utf-8 -*-
import sys, re, html, os; sys.path.insert(0,'.')
import _tpl, doc_privacy, doc_marketing, doc_offer
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = '/Users/glebastasev/клод/outputs'
os.makedirs(OUT, exist_ok=True)
DATE = '5 сентября 2026 года'

def blocks(src, keep_req=False):
    """Разбирает HTML документа на плоский список блоков."""
    if not keep_req:
        src = re.sub(r'(?s)<div class="req">.*?</div>', '', src)
    src = re.sub(r'(?s)<dl>.*?</dl>', lambda m: m.group(0), src)
    out = []
    for m in re.finditer(r'(?s)<(h2|h3|p|ul|dl|div)([^>]*)>(.*?)</\1>', src):
        tag, attrs, inner = m.group(1), m.group(2), m.group(3)
        if tag == 'div':
            for b in blocks(inner): out.append(b)
            continue
        if tag == 'ul':
            for li in re.findall(r'(?s)<li>(.*?)</li>', inner):
                out.append(('li', li))
        elif tag == 'dl':
            for dt, dd in re.findall(r'(?s)<dt>(.*?)</dt>\s*<dd>(.*?)</dd>', inner):
                out.append(('dl', f'{dt}: {dd}'))
        else:
            out.append((tag, inner))
    return out

def runs(par, frag):
    """Пишет фрагмент с учётом <b>, <a> и пропусков."""
    frag = re.sub(r'<a [^>]*>(.*?)</a>', r'\1', frag)
    frag = frag.replace('<span class="fill">', '<FILL>').replace('</span>', '</FILL>')
    for piece in re.split(r'(<b>.*?</b>|<strong>.*?</strong>|<FILL>.*?</FILL>)', frag, flags=re.S):
        if not piece: continue
        bold = piece.startswith(('<b>', '<strong>'))
        fill = piece.startswith('<FILL>')
        txt = html.unescape(re.sub(r'<[^>]+>', '', piece)).strip()
        if not txt: continue
        if not par.runs: pass
        elif not par.text.endswith((' ', '(', '«')) and not txt.startswith((' ', ',', '.', ';', ':', ')', '»')):
            txt = ' ' + txt
        r = par.add_run(('[' + txt.strip() + ']') if fill else txt)
        r.bold = bold or fill
        if fill: r.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)

def build(fname, title, sub, body):
    d = Document()
    st = d.styles['Normal']; st.font.name = 'Times New Roman'; st.font.size = Pt(11)
    st.paragraph_format.space_after = Pt(6); st.paragraph_format.line_spacing = 1.15
    for s in d.sections:
        s.top_margin = s.bottom_margin = Cm(2); s.left_margin = Cm(2.5); s.right_margin = Cm(1.5)
    h = d.add_paragraph(); h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = h.add_run(title.upper()); r.bold = True; r.font.size = Pt(14)
    p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(sub); r.italic = True; r.font.size = Pt(10)
    for tag, frag in blocks(body):
        if tag == 'h2':
            par = d.add_paragraph(); par.paragraph_format.space_before = Pt(14)
            runs(par, frag)
            for r in par.runs: r.bold = True; r.font.size = Pt(12)
        elif tag == 'h3':
            par = d.add_paragraph(); par.paragraph_format.space_before = Pt(10)
            runs(par, frag)
            for r in par.runs: r.bold = True
        elif tag == 'li':
            par = d.add_paragraph(style='List Bullet'); par.paragraph_format.space_after = Pt(2)
            runs(par, frag)
        elif tag == 'dl':
            par = d.add_paragraph(); par.paragraph_format.left_indent = Cm(0.6)
            par.paragraph_format.space_after = Pt(2); runs(par, frag)
        else:
            par = d.add_paragraph(); par.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            runs(par, frag)
    d.add_page_break()
    par = d.add_paragraph(); r = par.add_run('РЕКВИЗИТЫ ИСПОЛНИТЕЛЯ'); r.bold = True
    for tag, frag in blocks(_tpl.REQ, keep_req=True):
        if tag in ('dl',):
            par = d.add_paragraph(); par.paragraph_format.space_after = Pt(2); runs(par, frag)
    path = os.path.join(OUT, fname); d.save(path); print(path)

build('Оферта — ИП Заварзин.docx', 'Публичная оферта',
      'на заключение договора возмездного оказания информационно-консультационных услуг\nРедакция от ' + DATE, doc_offer.BODY)
build('Политика персональных данных — ИП Заварзин.docx', 'Политика обработки персональных данных',
      'Редакция от ' + DATE, doc_privacy.BODY)
build('Согласие на рассылки — ИП Заварзин.docx', 'Согласие на получение рекламных и информационных материалов',
      'Редакция от ' + DATE, doc_marketing.BODY)
